import os
import io
from typing import Union

try:
    import pdfplumber
    HAS_PDFPLUMBER = True
except ImportError:
    HAS_PDFPLUMBER = False

try:
    import pypdf
    HAS_PYPDF = True
except ImportError:
    HAS_PYPDF = False


def extract_text_from_file(file_source: Union[str, bytes, io.BytesIO], filename: str) -> str:
    ext = os.path.splitext(filename)[1].lower()

    if ext == ".txt":
        return _extract_txt(file_source)
    elif ext == ".pdf":
        return _extract_pdf(file_source)
    elif ext == ".docx":
        return _extract_docx(file_source)
    elif ext == ".md":
        return _extract_md(file_source)
    else:
        raise ValueError(f"Unsupported file format '{ext}'. Only .pdf, .txt, .docx, and .md files are supported.")


def _extract_txt(file_source: Union[str, bytes, io.BytesIO]) -> str:
    if isinstance(file_source, str):
        with open(file_source, "r", encoding="utf-8", errors="replace") as f:
            content = f.read()
    elif isinstance(file_source, bytes):
        content = file_source.decode("utf-8", errors="replace")
    elif hasattr(file_source, "read"):
        if hasattr(file_source, "seek"):
            file_source.seek(0)
        data = file_source.read()
        if isinstance(data, bytes):
            content = data.decode("utf-8", errors="replace")
        else:
            content = str(data)
    else:
        raise ValueError("Invalid file source type for TXT extraction.")

    content = content.strip()
    if not content:
        raise ValueError("The uploaded TXT document is empty.")
    return content


def _extract_pdf(file_source: Union[str, bytes, io.BytesIO]) -> str:
    text_pages = []

    if isinstance(file_source, str):
        stream = file_source
    elif isinstance(file_source, bytes):
        stream = io.BytesIO(file_source)
    else:
        stream = file_source

    if HAS_PDFPLUMBER:
        try:
            if hasattr(stream, "seek"):
                stream.seek(0)
            with pdfplumber.open(stream) as pdf:
                for page in pdf.pages:
                    extracted = page.extract_text()
                    if extracted:
                        text_pages.append(extracted.strip())
        except Exception:
            text_pages = []

    if not text_pages and HAS_PYPDF:
        try:
            if hasattr(stream, "seek"):
                stream.seek(0)
            reader = pypdf.PdfReader(stream)
            for page in reader.pages:
                extracted = page.extract_text()
                if extracted:
                    text_pages.append(extracted.strip())
        except Exception as e:
            raise ValueError(f"Failed to extract text from PDF: {str(e)}")

    full_text = "\n\n".join(text_pages).strip()
    if not full_text:
        raise ValueError(
            "Could not extract readable text from the PDF document. "
            "It may be empty, image-only (scanned), or password-protected."
        )
    return full_text


def _extract_docx(file_source: Union[str, bytes, io.BytesIO]) -> str:
    try:
        from docx import Document
    except ImportError:
        raise RuntimeError("python-docx is not installed, but is required for DOCX extraction.")

    if isinstance(file_source, str):
        doc = Document(file_source)
    elif isinstance(file_source, bytes):
        doc = Document(io.BytesIO(file_source))
    else:
        if hasattr(file_source, "seek"):
            file_source.seek(0)
        doc = Document(file_source)

    text_parts = []
    # Extract paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text.strip())

    # Extract tables
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                text_parts.append(" | ".join(row_text))

    full_text = "\n\n".join(text_parts).strip()
    if not full_text:
        raise ValueError("The uploaded DOCX document is empty or has no readable text.")
    return full_text


def _extract_md(file_source: Union[str, bytes, io.BytesIO]) -> str:
    # Markdown is parsed as raw text
    try:
        content = _extract_txt(file_source)
    except ValueError as e:
        if "TXT" in str(e):
            raise ValueError("The uploaded MD document is empty.")
        raise
    return content
